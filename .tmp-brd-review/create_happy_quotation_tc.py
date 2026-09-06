from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path(r"D:\Source\MySource\JogetMVC\TC_Happy_Path_Quotation.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9D9D9", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_run_font(run, name="Arial", size=9, bold=None, color="000000"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_text(cell, text, bold=False, color="000000", align=WD_ALIGN_PARAGRAPH.LEFT, size=8.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(10 if level == 1 else 7)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_run_font(run, size=14 if level == 1 else 11, bold=True)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.12
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, size=9.5, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, size=9.5)
    else:
        r = p.add_run(text)
        set_run_font(r, size=9.5)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.55)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_run_font(r, size=9.2)


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=0, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(3.2), Cm(6.0), Cm(3.2), Cm(14.0)]
    for i, row_data in enumerate(rows):
        cells = table.add_row().cells
        for j, width in enumerate(widths):
            cells[j].width = width
        set_cell_text(cells[0], row_data[0], bold=True, color="FFFFFF")
        set_cell_shading(cells[0], "1F4E78")
        set_cell_text(cells[1], row_data[1])
        set_cell_text(cells[2], row_data[2], bold=True, color="FFFFFF")
        set_cell_shading(cells[2], "1F4E78")
        set_cell_text(cells[3], row_data[3])
        if i % 2 == 1:
            set_cell_shading(cells[1], "F3F6F9")
            set_cell_shading(cells[3], "F3F6F9")
        keep_row(table.rows[-1])
    set_table_borders(table)
    return table


doc = Document()
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width, section.page_height = section.page_height, section.page_width
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(1.6)
section.right_margin = Cm(1.6)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
styles["Title"].font.name = "Arial"
styles["Title"].font.color.rgb = RGBColor(0, 0, 0)
for name in ("Heading 1", "Heading 2"):
    styles[name].font.name = "Arial"
    styles[name].font.color.rgb = RGBColor(0, 0, 0)

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.paragraph_format.space_after = Pt(4)
run = title.add_run("Test Case Happy Path Quotation")
set_run_font(run, size=22, bold=True)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(12)
r = subtitle.add_run("Báo giá tiêu chuẩn qua TS không referral UW và được khách hàng xác nhận")
set_run_font(r, size=11, color="404040")

add_body(
    doc,
    "Mục tiêu của test case là xác nhận một yêu cầu báo giá có thể đi xuyên suốt từ FO, TS, MKT Manager đến khách hàng và kết thúc ở trạng thái hoàn tất, không phát sinh return, revise, decline hoặc UW referral.",
)

add_heading(doc, "Thông tin test case")
add_key_value_table(doc, [
    ("Test Case ID", "QUO HP 001", "Mức ưu tiên", "High"),
    ("Module", "Quotation", "Loại kiểm thử", "Functional and workflow integration"),
    ("Workflow", "DEMO FLOW Quotation", "Kết quả cuối", "Quotation Confirmed and Completed"),
    ("Kịch bản", "New Business", "Phạm vi", "FO to TS to FO to LMKT to FO to Client Confirmed"),
    ("Ngoài phạm vi", "UW referral, return, revise, cancel, client refusal", "Trạng thái tài liệu", "Mẫu để review format"),
])

add_heading(doc, "Điều kiện tiên quyết")
for item in [
    "Các tài khoản FO, TS và LMKT đang hoạt động, được gán đúng role và có quyền truy cập quotation tương ứng.",
    "WorkflowDefinition DEMO FLOW đang active; các action trong kịch bản tồn tại và active trong StepsWorkflow.",
    "Khách hàng và dữ liệu master cần thiết như Line of Business, sản phẩm và tiền tệ đã tồn tại.",
    "Mẫu quotation, email template và notification template cần dùng đã được cấu hình.",
    "Tệp đính kèm hợp lệ, không vượt giới hạn dung lượng và đúng định dạng hệ thống cho phép.",
]:
    add_bullet(doc, item)

add_heading(doc, "Dữ liệu kiểm thử")
data_table = doc.add_table(rows=1, cols=4)
data_table.alignment = WD_TABLE_ALIGNMENT.CENTER
data_table.autofit = False
for i, width in enumerate([Cm(4.2), Cm(9.0), Cm(4.2), Cm(9.0)]):
    data_table.rows[0].cells[i].width = width
headers = ["Trường", "Giá trị mẫu", "Trường", "Giá trị mẫu"]
for i, value in enumerate(headers):
    set_cell_text(data_table.rows[0].cells[i], value, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(data_table.rows[0].cells[i], "1F4E78")
repeat_header(data_table.rows[0])
test_data = [
    ("Transaction Type", "New Business", "Refer UW", "No"),
    ("Client", "TC Quotation Happy Client", "Line of Business", "Fire"),
    ("Effective Date", "Ngày hợp lệ trong tương lai", "Expiry Date", "Effective Date cộng 1 năm"),
    ("Số quotation", "1", "Quotation Code", "Hệ thống sinh Q1"),
    ("Attachment đầu vào", "request detail.pdf", "Quotation document", "quotation Q1.pdf"),
    ("Client confirmation", "Accepted", "Confirmation attachment", "client confirmation.pdf"),
]
for idx, row_data in enumerate(test_data):
    cells = data_table.add_row().cells
    for i, value in enumerate(row_data):
        set_cell_text(cells[i], value, bold=(i % 2 == 0))
        if idx % 2 == 1:
            set_cell_shading(cells[i], "F3F6F9")
    keep_row(data_table.rows[-1])
set_table_borders(data_table)

add_heading(doc, "Các bước thực hiện")
steps = [
    ("1", "FO", "Đăng nhập, mở chức năng tạo Quotation và chọn New Business.", "Màn hình tạo mới hiển thị đúng quyền FO; các trường bắt buộc và vùng attachment sẵn sàng nhập.", "Draft"),
    ("2", "FO", "Nhập đầy đủ dữ liệu khách hàng, sản phẩm, thời hạn bảo hiểm; đính kèm request detail.pdf và lưu.", "Hệ thống tạo quotation record duy nhất, sinh mã tham chiếu, lưu dữ liệu và attachment; không tạo workflow trùng.", "FO Process"),
    ("3", "FO", "Chọn xử lý qua TS và thực hiện Submit To TS.", "Action SUBMIT MAIN TS thành công; task chuyển cho TS, status là TS Pending và người liên quan nhận notification.", "TS Pending"),
    ("4", "TS", "Mở task từ inbox; kiểm tra dữ liệu, nhập thông tin tính phí, chọn số lượng quotation là 1 và tạo quotation Q1.", "TS xem được toàn bộ dữ liệu cần thiết; hệ thống tạo đúng một quotation code Q1 và lưu quotation document tương ứng.", "TS Process"),
    ("5", "TS", "Đính kèm quotation Q1.pdf và thực hiện Submit To FO.", "Action SUBMIT TS FO thành công; task trở về FO, status là FO Process; file và dữ liệu TS nhập không bị mất.", "FO Process"),
    ("6", "FO", "Review quotation; xác nhận không cần referral UW và gửi Asking Signature Approval đến LMKT.", "Action SUBMIT FO LMKT thành công; task chuyển LMKT, status là MKT Review; không tạo task UW.", "MKT Review"),
    ("7", "LMKT", "Mở task, kiểm tra quotation và chọn Approved Return to FO; hoàn tất chữ ký theo cấu hình.", "Action LMKT APPROVED FO thành công; bản ký được lưu đúng version, task quay lại FO và có notification.", "FO Review"),
    ("8", "FO", "Kiểm tra bản đã duyệt và thực hiện Send Quotation to Client.", "Action SEND QUOTATION CLIENT thành công; email dùng đúng người nhận, subject và attachment; status chuyển Waiting Client Confirmation.", "Waiting Client Confirmation"),
    ("9", "FO", "Sau khi khách hàng đồng ý, chọn Client Confirmed và tải lên client confirmation.pdf.", "Action CLIENT CONFIRMED thành công; xác nhận và attachment được lưu; không còn task xử lý mở.", "Quotation Confirmed"),
    ("10", "FO", "Mở lại quotation, audit log, dashboard và danh sách tìm kiếm.", "Quotation ở trạng thái Completed; lịch sử đủ actor, action và thời gian; tìm thấy theo quotation code, client và PIC; không thể chạy lại action kết thúc.", "Completed"),
]
steps_table = doc.add_table(rows=1, cols=5)
steps_table.alignment = WD_TABLE_ALIGNMENT.CENTER
steps_table.autofit = False
widths = [Cm(1.1), Cm(1.7), Cm(8.2), Cm(10.6), Cm(4.6)]
headers = ["STT", "Role", "Thao tác", "Kết quả mong đợi", "Status mong đợi"]
for i, (cell, width, header) in enumerate(zip(steps_table.rows[0].cells, widths, headers)):
    cell.width = width
    set_cell_text(cell, header, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER, size=8.5)
    set_cell_shading(cell, "1F4E78")
repeat_header(steps_table.rows[0])
for idx, row_data in enumerate(steps):
    cells = steps_table.add_row().cells
    for i, (cell, width, value) in enumerate(zip(cells, widths, row_data)):
        cell.width = width
        align = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 1, 4) else WD_ALIGN_PARAGRAPH.LEFT
        set_cell_text(cell, value.replace(" ", "_") if False else value, align=align, size=8.1)
        if idx % 2 == 1:
            set_cell_shading(cell, "F3F6F9")
    keep_row(steps_table.rows[-1])
set_table_borders(steps_table)

add_heading(doc, "Đối chiếu workflow và database")
add_body(doc, "Các action code kỳ vọng theo thứ tự:")
action_p = doc.add_paragraph()
action_p.paragraph_format.space_after = Pt(6)
action_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for i, action in enumerate([
    "SUBMIT_MAIN_TS", "SUBMIT_TS_FO", "SUBMIT_FO_LMKT", "LMKT_APPROVED_FO", "SEND_QUOTATION_CLIENT", "CLIENT_CONFIRMED"
]):
    if i:
        sep = action_p.add_run("  →  ")
        set_run_font(sep, size=9, color="666666")
    rr = action_p.add_run(action)
    set_run_font(rr, size=9, bold=True, color="1F4E78")

add_body(doc, "Chỉ truy vấn các cột cần thiết để giảm dữ liệu trả về. Thay biến @QuotationGuid bằng Guid của quotation vừa test.")

queries = [
    ("Kiểm tra instance cuối workflow", "SELECT TOP 1 Id, Guid, RecordGuid, WorkflowDefinitionId, CurrentStep, CurrentOwnerRoleCode, LastActionCode, StartedDate, CompletedDate, IsCompleted, IsCancelled\nFROM InstanceWorkflow\nWHERE RecordGuid = @QuotationGuid\nORDER BY Id DESC;"),
    ("Kiểm tra action cấu hình", "SELECT Id, ActionCode, FromNodeId, ToNodeId, StatusId, StatusName, IsReturn, IsEnd, IsActive\nFROM StepsWorkflow\nWHERE WorkflowDefinitionId = @WorkflowDefinitionId\n  AND ActionCode IN ('SUBMIT_MAIN_TS','SUBMIT_TS_FO','SUBMIT_FO_LMKT','LMKT_APPROVED_FO','SEND_QUOTATION_CLIENT','CLIENT_CONFIRMED')\nORDER BY SortOrder;"),
    ("Kiểm tra CurrentStep thuộc definition", "SELECT Id, Code, WorkflowDefinitionId, IsActive, LastActionCode\nFROM WorkflowInstanceNode\nWHERE WorkflowDefinitionId = @WorkflowDefinitionId\n  AND Code = @CurrentStep;"),
]
for title_text, sql in queries:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title_text)
    set_run_font(r, size=9.2, bold=True)
    sql_table = doc.add_table(rows=1, cols=1)
    sql_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = sql_table.cell(0, 0)
    set_cell_shading(cell, "F3F6F9")
    set_cell_text(cell, sql, size=8.1)
    for run in cell.paragraphs[0].runs:
        set_run_font(run, name="Consolas", size=8.1)
    set_table_borders(sql_table, color="D9D9D9")

add_heading(doc, "Tiêu chí Pass")
for item in [
    "Một quotation duy nhất được tạo và đi đúng thứ tự FO to TS to FO to LMKT to FO to Client.",
    "Không phát sinh UW task vì Refer UW bằng No.",
    "Dữ liệu, version quotation và attachment được giữ nguyên qua mọi lần chuyển bước.",
    "Notification, email và audit log ghi nhận đúng actor, action, người nhận và thời điểm.",
    "Kết thúc với LastActionCode CLIENT CONFIRMED, CompletedDate có giá trị, IsCompleted bằng 1 và IsCancelled bằng 0.",
]:
    add_bullet(doc, item)

add_heading(doc, "Điểm cần ghi nhận khi chạy thực tế")
add_body(doc, "Nếu giao diện hiển thị Completed nhưng InstanceWorkflow vẫn có IsCompleted bằng 0, CompletedDate trống hoặc action kết thúc không khóa workflow, ghi nhận defect về đồng bộ trạng thái kết thúc. Đây là điểm đã được nhận diện khi đối chiếu dữ liệu hiện tại với yêu cầu BRD.")

footer = section.footer
p = footer.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Workflow Management  |  QUO HP 001")
set_run_font(r, size=8, color="666666")

doc.core_properties.title = "Test Case Happy Path Quotation"
doc.core_properties.subject = "Quotation workflow happy path test case"
doc.core_properties.author = "QA Team"
doc.save(OUT)
print(OUT)
