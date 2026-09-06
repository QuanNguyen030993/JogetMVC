from pathlib import Path
from zipfile import ZipFile
import json
import re

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


SOURCE = Path(r"D:\Folder drive\BRD_Workflow Managermentt v1.0 3 updated.docx")
OUT = Path(r"D:\Source\MySource\JogetMVC\.tmp-brd-review")


def iter_blocks(parent):
    body = parent.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def clean(value):
    return re.sub(r"[ \t]+", " ", value or "").strip()


doc = Document(SOURCE)
lines = []
table_count = 0

for block in iter_blocks(doc):
    if isinstance(block, Paragraph):
        text = clean(block.text)
        if not text:
            continue
        style = block.style.name if block.style else ""
        if style.lower().startswith("heading") or style.lower() == "title":
            lines.append(f"\n## [{style}] {text}")
        else:
            lines.append(text)
    else:
        table_count += 1
        lines.append(f"\n### [TABLE {table_count}]")
        for row_index, row in enumerate(block.rows, 1):
            cells = [clean(cell.text).replace("\n", " / ") for cell in row.cells]
            lines.append(f"ROW {row_index}: " + " | ".join(cells))

for section_index, section in enumerate(doc.sections, 1):
    header = " | ".join(clean(p.text) for p in section.header.paragraphs if clean(p.text))
    footer = " | ".join(clean(p.text) for p in section.footer.paragraphs if clean(p.text))
    if header:
        lines.append(f"\n### [HEADER SECTION {section_index}] {header}")
    if footer:
        lines.append(f"\n### [FOOTER SECTION {section_index}] {footer}")

(OUT / "content.txt").write_text("\n".join(lines), encoding="utf-8")

with ZipFile(SOURCE) as zf:
    names = zf.namelist()
    media = [name for name in names if name.startswith("word/media/") and not name.endswith("/")]
    comments = [name for name in names if name == "word/comments.xml"]
    document_xml = zf.read("word/document.xml").decode("utf-8", errors="replace")
    image_dir = OUT / "media"
    image_dir.mkdir(exist_ok=True)
    for name in media:
        (image_dir / Path(name).name).write_bytes(zf.read(name))

summary = {
    "paragraphs": len(doc.paragraphs),
    "tables": len(doc.tables),
    "sections": len(doc.sections),
    "inline_shapes": len(doc.inline_shapes),
    "media_files": len(media),
    "has_comments": bool(comments),
    "tracked_insertions": document_xml.count("<w:ins"),
    "tracked_deletions": document_xml.count("<w:del"),
    "manual_page_breaks": document_xml.count('w:type="page"'),
    "content_lines": len(lines),
}
(OUT / "summary.json").write_text(json.dumps(summary, indent=2), encoding="utf-8")
print(json.dumps(summary, ensure_ascii=False))
